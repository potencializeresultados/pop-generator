import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
from sqlalchemy.orm import Session
from app.models import POP, Company
import re

# Caminho para o template base
TEMPLATE_PATH = "templates/POP_base.docx"

def hex_to_rgb(hex_color):
    """Converte cor hexadecimal para RGB"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def process_description(description):
    """
    Processa a descrição do POP conforme as regras:
    - Cada etapa deve começar com verbo no infinitivo
    - Encerrar cada passo com ";" (ponto-vírgula)
    - Após o último passo, inserir "Fim do Processo"
    """
    # Dividir a descrição em linhas
    steps = description.split('\n')
    
    # Processar cada passo
    processed_steps = []
    for step in steps:
        step = step.strip()
        if not step:
            continue
            
        # Verificar se começa com verbo no infinitivo (simplificado)
        if not re.match(r'^[A-Za-zÀ-ÖØ-öø-ÿ]+ar\b|^[A-Za-zÀ-ÖØ-öø-ÿ]+er\b|^[A-Za-zÀ-ÖØ-öø-ÿ]+ir\b', step):
            step = "Realizar " + step
            
        # Garantir que termina com ponto e vírgula
        if not step.endswith(';'):
            step += ';'
            
        processed_steps.append(step)
    
    # Adicionar "Fim do Processo" após o último passo
    if processed_steps:
        processed_steps.append("Fim do processo.")
    
    return '\n'.join(processed_steps)

def create_cover_page(doc, company, pop):
    """Cria uma capa personalizada para o POP"""
    # Adicionar uma nova seção para a capa
    section = doc.add_section()
    
    # Definir margens da página
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # Converter cor primária de hex para RGB
    try:
        r, g, b = hex_to_rgb(company.primary_color)
    except:
        r, g, b = (255, 0, 0)  # Vermelho padrão se a conversão falhar
    
    # Adicionar título da empresa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company.name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(r, g, b)
    
    # Adicionar espaço
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Adicionar logotipo (se existir)
    if os.path.exists(company.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(company.logo_path, width=Inches(3))
    
    # Adicionar espaço
    for _ in range(3):
        doc.add_paragraph()
    
    # Adicionar título do POP
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{pop.pop_code} - {pop.department}")
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Adicionar nome do procedimento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pop.procedure_name)
    run.font.size = Pt(20)
    run.font.bold = True
    
    # Adicionar quebra de página
    doc.add_page_break()
    
    return doc

def generate_pop_document(pop_id, company, output_path, db: Session):
    """
    Gera o documento DOCX do POP com base no template
    """
    # Obter dados do POP
    pop = db.query(POP).filter(POP.id == pop_id).first()
    if not pop:
        raise ValueError("POP não encontrado")
    
    # Processar a descrição
    processed_description = process_description(pop.description)
    
    # Criar documento a partir do template
    doc = Document(TEMPLATE_PATH)
    
    # Criar capa personalizada
    doc = create_cover_page(doc, company, pop)
    
    # Modificar o documento para substituir [RIBAS] pelo nome da empresa
    for paragraph in doc.paragraphs:
        if "[RIBAS]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[RIBAS]", company.name)
    
    # Atualizar os metadados do POP na tabela
    # Assumindo que a primeira tabela contém os metadados
    if doc.tables:
        table = doc.tables[0]
        
        # Mapeamento de células da tabela para valores do POP
        # Ajuste os índices conforme a estrutura real da tabela no template
        cell_mappings = {
            (0, 1): pop.procedure_name,  # Nome do Procedimento
            (1, 1): pop.pop_code,        # Código POP
            (2, 1): pop.department,      # Departamento
            (3, 1): pop.periodicity,     # Periodicidade
            (4, 1): pop.start_date,      # Data Início
            (5, 1): pop.target_date or "N/A",  # Data Meta
            (6, 1): pop.legal_date or "N/A",   # Data Legal
            (7, 1): pop.competence,      # Competência
            (8, 1): pop.tax_regime,      # Regime Tributário
            (9, 1): pop.complexity,      # Complexidade
            (10, 1): pop.responsible,    # Responsável
            (11, 1): pop.automatic_routine,  # Rotina Automática/Lote
            (12, 1): pop.estimated_time, # Tempo Médio Estimado
        }
        
        # Atualizar células da tabela
        for (row, col), value in cell_mappings.items():
            if row < len(table.rows) and col < len(table.rows[row].cells):
                table.rows[row].cells[col].text = value
    
    # Adicionar a descrição processada
    # Encontrar o parágrafo que contém "DESCRIÇÃO DO PROCEDIMENTO"
    description_added = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "DESCRIÇÃO DO PROCEDIMENTO" in paragraph.text:
            # Adicionar a descrição após este parágrafo
            for step in processed_description.split('\n'):
                if step:
                    p = doc.add_paragraph(step, style='Normal')
                    p.paragraph_format.space_after = Pt(6)
            description_added = True
            break
    
    # Se não encontrou o parágrafo específico, adiciona ao final
    if not description_added:
        doc.add_heading("DESCRIÇÃO DO PROCEDIMENTO:", level=1)
        for step in processed_description.split('\n'):
            if step:
                p = doc.add_paragraph(step, style='Normal')
                p.paragraph_format.space_after = Pt(6)
    
    # Salvar o documento
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    return output_path

def generate_cover_preview(company_id, db: Session):
    """
    Gera uma prévia da capa do POP
    """
    # Obter dados da empresa
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise ValueError("Empresa não encontrada")
    
    # Converter cor primária de hex para RGB
    try:
        r, g, b = hex_to_rgb(company.primary_color)
    except:
        r, g, b = (255, 0, 0)  # Vermelho padrão se a conversão falhar
    
    # Criar um documento em branco
    doc = Document()
    
    # Configurar margens
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # Adicionar título da empresa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company.name)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(r, g, b)
    
    # Adicionar espaço
    doc.add_paragraph()
    
    # Adicionar logotipo (se existir)
    if os.path.exists(company.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(company.logo_path, width=Inches(3))
    
    # Adicionar espaço
    for _ in range(2):
        doc.add_paragraph()
    
    # Adicionar título de exemplo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CÓDIGO.XX - DEPARTAMENTO")
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Adicionar nome do procedimento de exemplo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOME DO PROCEDIMENTO")
    run.font.size = Pt(20)
    run.font.bold = True
    
    # Adicionar informações sobre as cores
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRÉVIA DA CAPA")
    run.font.size = Pt(14)
    run.font.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Cor Primária: {company.primary_color}")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Cor Secundária: {company.secondary_color}")
    run.font.size = Pt(12)
    
    # Salvar prévia temporária
    preview_path = f"uploads/previews/cover_preview_{company_id}.docx"
    os.makedirs("uploads/previews", exist_ok=True)
    doc.save(preview_path)
    
    return preview_path
